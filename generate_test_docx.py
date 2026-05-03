"""Generate a comprehensive test DOCX with all cases we discussed.
Uses REAL paragraph structure (not \n in single paragraph).
"""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Option B Comprehensive Test Cases', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('This document tests all matching scenarios for putpnginword.py Option B (block-scoped node matching).')
doc.add_paragraph()

# Helper to add a test case with REAL paragraphs
def add_test_case(doc, case_num, case_name, description, cell_lines):
    """cell_lines: list of strings, each becomes a separate paragraph"""
    doc.add_heading(f'Case {case_num}: {case_name}', level=1)
    doc.add_paragraph(f'Description: {description}')
    
    # Create table with 2 rows (header + content)
    table = doc.add_table(rows=2, cols=2)
    table.style = 'Table Grid'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Test Case'
    header_cells[1].text = 'Cell Content'
    
    # Content
    content_cells = table.rows[1].cells
    content_cells[0].text = f'Case {case_num}'
    
    # Add each line as a SEPARATE paragraph in the cell
    cell = content_cells[1]
    cell.text = ''  # Clear default paragraph
    for line in cell_lines:
        para = cell.add_paragraph(line)
    
    doc.add_paragraph()

# Case 1: Basic single block
add_test_case(
    doc, 1, 'Basic Single Block',
    'Simple command with multiple nodes. Should insert 3 images (1 per node).',
    [
        '<HUAWEI>display device',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
    ]
)

# Case 2: Nested command block
add_test_case(
    doc, 2, 'Nested Command Block',
    'Nested commands (system-view + subview) with quit. Should insert 1 image.',
    [
        '<HUAWEI>system-view',
        '[HW]interface GigabitEthernet0/0/1',
        '[HW-GigabitEthernet0/0/1]display this',
        '[HW-GigabitEthernet0/0/1]quit',
        '[HW]quit',
        '<TUC-TEST01>',
    ]
)

# Case 3: Multiple blocks per cell (Option B - KEY CASE)
add_test_case(
    doc, 3, 'Multiple Blocks Per Cell (Option B)',
    'Two separate command blocks each followed by same nodes. Should insert 6 images (2 blocks x 3 nodes).',
    [
        '<HUAWEI>display device',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
        '<HUAWEI>display clock',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
    ]
)

# Case 4: Multi-model scenario
add_test_case(
    doc, 4, 'Multi-Model Scenario',
    'Different NEs with devices appearing after specific blocks. device2 matches display device (block 0), device1 matches display clock (block 1).',
    [
        '<TUC-TEST01>display device',
        '<TUC-TEST02>',
        '<TUC-TEST01>display clock',
        '<TUC-TEST01>',
    ]
)

# Case 5: Error preference
add_test_case(
    doc, 5, 'Error Preference',
    'Command that produces [error] PNG should be skipped. No image should be inserted.',
    [
        '<TUC-TEST01>display cpu-usage',
        '<TUC-TEST01>',
    ]
)

# Case 6: Block isolation
add_test_case(
    doc, 6, 'Block Isolation',
    'Same node appears after two different commands. Should get 2 images (display device + display clock).',
    [
        '<HUAWEI>display device',
        '<TUC-TEST01>',
        '<HUAWEI>display clock',
        '<TUC-TEST01>',
    ]
)

# Case 7: Cell without quit
add_test_case(
    doc, 7, 'Cell Without Quit',
    'Nested block without quit should still match. Should insert 1 image.',
    [
        '<HUAWEI>system-view',
        '[HW]interface GigabitEthernet0/0/1',
        '[HW-GigabitEthernet0/0/1]display this',
        '<TUC-TEST01>',
    ]
)

# Case 8: Empty nodes (no nodes)
add_test_case(
    doc, 8, 'Cell With No Nodes',
    'Cell with command but no nodes. Should insert 0 images.',
    [
        '<HUAWEI>display device',
    ]
)

# Case 9: Wrong command
add_test_case(
    doc, 9, 'Wrong Command - No PNG',
    'Command with no matching PNG. Should insert 0 images.',
    [
        '<HUAWEI>display wrong-command',
        '<TUC-TEST01>',
    ]
)

# Case 10: Real multi-block (like Row 7)
add_test_case(
    doc, 10, 'Real Multi-Block (3 blocks)',
    'Three command blocks each followed by 3 nodes. Should insert 9 images (3 blocks x 3 nodes).',
    [
        '<HUAWEI>system-view',
        '[HW]interface GigabitEthernet0/0/1',
        '[HW-GigabitEthernet0/0/1]display this',
        '[HW-GigabitEthernet0/0/1]quit',
        '[HW]quit',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
        '<HUAWEI>display device',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
        '<HUAWEI>display clock',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
    ]
)

# Case 11: Error + OK blocks
add_test_case(
    doc, 11, 'Error Block + OK Block',
    'First block has wrong command (no match), second block has valid command. Should insert 3 images for second block only.',
    [
        '<HUAWEI>display wrong-command',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
        '<HUAWEI>display device',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
    ]
)

# Case 12: Multi-model with empty blocks (the key test)
add_test_case(
    doc, 12, 'Multi-Model Empty Block Merge',
    'CEx01 model supports display device, CEx02 supports display clock. Empty block should merge so all devices try both commands.',
    [
        'CEx01',
        '<HUAWEI>display device',
        'CEx02',
        '<HUAWEI>display clock',
        '<TUC-TEST01>',
        '<TUC-TEST02>',
        '<TUC-TEST03>',
    ]
)

# Save
doc.save('comprehensive_test_cases.docx')
print("Generated: comprehensive_test_cases.docx")
print("\nExpected image counts per case:")
print("  Case 1: 3 images")
print("  Case 2: 1 image")
print("  Case 3: 6 images (Option B)")
print("  Case 4: 2 images (multi-model)")
print("  Case 5: 0 images (error)")
print("  Case 6: 2 images")
print("  Case 7: 1 image")
print("  Case 8: 0 images")
print("  Case 9: 0 images")
print("  Case 10: 9 images")
print("  Case 11: 3 images")
print("  Case 12: 2 images (merged empty block, devices try both commands)")
print("\nTotal expected: 28 images")
