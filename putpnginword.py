import re
import os
import glob
import shutil
import uuid
import json
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from word_com_embed import embed_txt_via_word
from filename_utils import sanitize_filename

NOISE_COMMANDS = {'username'}


def get_base_dir():
    """Get the directory where the .exe or .py script lives."""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


# --- Configuration ---
CONFIG_PATH = os.path.join(get_base_dir(), "putpnginword_config.json")

_DEFAULT_CONFIG = {
    "docx_input": "comprehensive_test_cases_updated.docx",
    "png_path": r"screenshots\*.png",
    "docx_output": "comprehensive_test_cases_updated_RESULT_v2.docx",
    "target_sections": []
}


# Removed hardcoded constants — now loaded from putpnginword_config.json:
# DOCX_INPUT, PNG_PATH, DOCX_OUTPUT, TARGET_SECTIONS


def _ensure_config():
    """If config.json is missing, create a template and exit."""
    if os.path.exists(CONFIG_PATH):
        return
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(_DEFAULT_CONFIG, f, indent=2, ensure_ascii=False)
    print(f"Created template: {CONFIG_PATH}")
    print("Please edit it with your paths, then run again.")
    sys.exit(0)


def load_config():
    """Read putpnginword_config.json and return settings."""
    _ensure_config()
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        cfg = json.load(f)
    return cfg


# --- Regex patterns ---
# Match prompt+command: <Router>cmd or [~*Router]cmd or [~*Router-sub]cmd
PROMPT_LINE_RE = re.compile(r'^(<[A-Za-z][\w.\-]*>|\[~?\*?[A-Za-z][\w.\-/]*\])\s*(.+)$')

# Match node-only: <NodeName> with nothing after
NODE_LINE_RE = re.compile(r'^<([A-Za-z][\w.\-]*)>\s*$')


def extract_title(text):
    """Remove ID prefix like T01-0201-01 from heading text."""
    m = re.match(r'^[A-Z]\d+(?:[.-]\d+)*\s+(.*)$', text)
    if m:
        return m.group(1).strip()
    return text.strip()


def get_heading_level(para):
    """Extract level number from 'Heading N' style."""
    if para.style and para.style.name:
        name = para.style.name
        if name.startswith('Heading'):
            parts = name.split()
            if len(parts) == 2 and parts[1].isdigit():
                return int(parts[1])
    return None


def should_process_table(document, table_idx, target_titles):
    """Walk body elements in order. Activate on heading match, deactivate on same/higher level."""
    para_map = {para._p: para for para in document.paragraphs}
    seen_tables = 0
    active_level = None

    for child in document.element.body:
        if child.tag.endswith('tbl'):
            if seen_tables == table_idx:
                return active_level is not None
            seen_tables += 1
            continue

        if child.tag.endswith('p'):
            para = para_map.get(child)
            if not para:
                continue

            level = get_heading_level(para)
            if level is None:
                continue

            title = extract_title(para.text)

            if title in target_titles:
                active_level = level
            elif active_level is not None and level <= active_level:
                active_level = None

    return False


def list_sections(doc):
    """Print all heading sections found in the document, in order."""
    print("Sections found in document:")
    for para in doc.paragraphs:
        level = get_heading_level(para)
        if level is not None and para.text.strip():
            title = extract_title(para.text)
            print(f"  [Heading {level}] {title}")


# --- Abbreviation Expansion ---

# Longest-first to avoid partial expansion (e.g. 'dis th' before 'dis')
_ABBREVIATIONS = [
    ('dis th', 'display this'),
    ('dis cur', 'display current-configuration'),
    ('dis cu', 'display current-configuration'),
    ('dis', 'display'),
    ('system', 'system-view'),
    ('sys', 'system-view'),
    ('comm', 'commit'),
    ('q', 'quit'),
]


def expand_abbreviations(commands: list[str]) -> list[str]:
    """Expand Huawei CLI abbreviations in commands.

    Uses longest-match-first to avoid partial expansion.
    Example: 'dis th' -> 'display this', not 'display th'.
    """
    result = []
    for cmd in commands:
        expanded = cmd
        for abbrev, full in _ABBREVIATIONS:
            # Match as whole words (or start of string) to avoid mid-word replacement
            pattern = re.compile(r'^(\s*)' + re.escape(abbrev) + r'(\s|$)')
            expanded = pattern.sub(r'\1' + full + r'\2', expanded)
        result.append(expanded)
    return result


def parse_paragraphs(paragraphs):
    """Extract command blocks and node names from cell paragraphs.

    Splits commands into separate blocks when a standalone user-view
    command (<Router>cmd) is encountered after a nested block has started.

    Returns list of (commands, nodes) tuples, one per logical block.
    """
    blocks = []
    current_commands = []
    nodes = []

    # Match prompt-only lines like [Router] or <Router> with no command
    PROMPT_ONLY_RE = re.compile(
        r'^(<[A-Za-z][\w.\-]*>|\[~?\*?[A-Za-z][\w.\-/]*\])\s*$'
    )

    for para in paragraphs:
        # Split paragraph text by newlines to handle Word table cells
        # where multiple lines are in a single paragraph
        for line in para.text.split('\n'):
            text = line.strip()
            if not text:
                continue

            # Check if this line is a node-only: <NodeName> with no command
            m = NODE_LINE_RE.match(text)
            if m:
                nodes.append(m.group(1))
                continue

            # Skip prompt-only lines (no command after prompt)
            if PROMPT_ONLY_RE.match(text):
                continue

            # Check if this line is a prompt+command
            m = PROMPT_LINE_RE.match(text)
            if m:
                cmd = m.group(2).strip()
                prompt = m.group(1)
                if not cmd:
                    continue

                # Detect standalone user-view command: <Router>cmd (depth 0)
                # If we already have commands in current block, this starts a new block
                if (prompt.startswith('<') and
                        not prompt.startswith('<~') and
                        not prompt.startswith('<*') and
                        current_commands and
                        cmd.split()[0].lower() not in NOISE_COMMANDS):
                    blocks.append((current_commands, nodes))
                    current_commands = []
                    nodes = []

                current_commands.append(cmd)

    if current_commands:
        blocks.append((current_commands, nodes))

    return blocks


def parse_paragraphs_detailed(paragraphs):
    """Extract command blocks and node names from cell paragraphs.

    Same logic as parse_paragraphs(), but returns paragraph indices for each node.
    Also detects "Error:" in output text after a command to mark blocks that expect
    an [error] PNG. Noise commands like "username" are handled in block matching, not filtered here.

    Returns list of (commands, [(node, para_idx), ...], expect_error) tuples.
    """
    blocks = []
    current_commands = []
    nodes = []  # [(node_name, para_idx), ...]
    expect_error = False

    PROMPT_ONLY_RE = re.compile(
        r'^(<[A-Za-z][\w.\-]*>|\[~?\*?[A-Za-z][\w.\-/]*\])\s*$'
    )
    ERROR_LINE_RE = re.compile(r'^Error:', re.IGNORECASE)

    def _flush_block():
        nonlocal current_commands, nodes, expect_error
        if current_commands or nodes:
            blocks.append((current_commands[:], nodes[:], expect_error))
            current_commands = []
            nodes = []
            expect_error = False

    for para_idx, para in enumerate(paragraphs):
        lines = para.text.split('\n')
        for line_idx, line in enumerate(lines):
            text = line.strip()
            if not text:
                continue

            m = NODE_LINE_RE.match(text)
            if m:
                nodes.append((m.group(1), para_idx))
                continue

            if PROMPT_ONLY_RE.match(text):
                continue

            m = PROMPT_LINE_RE.match(text)
            if m:
                cmd = m.group(2).strip()
                prompt = m.group(1)
                if not cmd:
                    continue

                # Check for standalone user-view command starting a new block
                if (prompt.startswith('<') and
                        not prompt.startswith('<~') and
                        not prompt.startswith('<*') and
                        current_commands and
                        cmd.split()[0].lower() not in NOISE_COMMANDS):
                    _flush_block()

                current_commands.append(cmd)

                # After a command, scan remaining lines in this paragraph
                # for "Error:" until next prompt/node
                for future_line in lines[line_idx + 1:]:
                    fl = future_line.strip()
                    if not fl:
                        continue
                    if PROMPT_LINE_RE.match(fl) or NODE_LINE_RE.match(fl):
                        break
                    if ERROR_LINE_RE.match(fl):
                        expect_error = True
                        break
                continue

            # Non-prompt, non-node line (output text)
            if ERROR_LINE_RE.match(text):
                expect_error = True

    if current_commands or nodes:
        blocks.append((current_commands, nodes, expect_error))

    return blocks


def _merge_empty_blocks(blocks):
    """Merge empty standalone blocks into the next block with nodes.
    
    If a block has commands but NO nodes, prepend its commands to the NEXT
    block that has nodes. Blocks that already have nodes are untouched.
    Propagates expect_error: if any merged block had expect_error=True,
    the merged result has expect_error=True.
    """
    if not blocks:
        return blocks

    merged = []
    pending_commands = []
    pending_error = False

    for bi, (commands, nodes, expect_error) in enumerate(blocks):
        if not commands and not nodes:
            continue

        if nodes:
            if pending_commands:
                merged_commands = pending_commands + commands
                merged_error = pending_error or expect_error
                merged.append((merged_commands, nodes, merged_error))
                pending_commands = []
                pending_error = False
            else:
                merged.append((commands, nodes, expect_error))
        elif commands:
            # commands but no nodes → accumulate for next block with nodes
            pending_commands.extend(commands)
            pending_error = pending_error or expect_error

    return merged


def match_cell_blocks(paragraphs, png_files):
    """Match PNGs to command blocks for a single cell (Option B + empty block merge).

    Empty standalone blocks (no nodes) are merged into the next block with nodes.
    For merged blocks, each node tries each command individually until finding
    a match. When expect_error=True for a block, prefers [error] PNGs.
    Returns list of dicts: {
        'node': str,
        'block_idx': int,
        'para_idx': int,
        'commands': list[str],
        'match_path': str | None,
        'action': 'inserted' | 'skipped_error' | 'no_match',
    }
    """
    blocks = parse_paragraphs_detailed(paragraphs)
    
    # Check if any merge happened
    original_blocks = blocks[:]
    blocks = _merge_empty_blocks(blocks)
    was_merged = len(blocks) < len(original_blocks)
    
    results = []

    for block_idx, (commands, block_nodes, expect_error) in enumerate(blocks):
        if not commands or not block_nodes:
            continue

        expanded_commands = expand_abbreviations(commands)

        # ── Strip leading SSH/telnet commands from matching ──
        # When the first command is stelnet/ssh/telnet and a display/dis
        # command exists later, use only the display command for filename matching.
        # This handles merged SSH sessions from process_network_logs.py.
        match_commands = expanded_commands
        if (expanded_commands and
            expanded_commands[0].strip().lower().startswith(('stelnet', 'ssh', 'telnet'))):
            for i, cmd in enumerate(expanded_commands):
                if cmd.strip().lower().startswith(('dis', 'display')):
                    match_commands = [cmd]
                    break

        for node, para_idx in block_nodes:
            png_match = None
            best_commands = None
            
            if was_merged:
                # For merged blocks, try each command individually until finding
                # a non-error match (or until we've tried all commands).
                for i in range(len(match_commands)):
                    single_cmd = match_commands[i]
                    match = find_best_match(
                        node, [single_cmd], png_files, prefer_error=expect_error
                    )
                    if not match:
                        continue

                    # Check if the matched PNG is an error PNG and whether we
                    # should skip it or accept it.
                    has_error = '[error]' in os.path.basename(match).lower()
                    if has_error and not expect_error:
                        # Error PNG but block doesn't expect it → skip and try next
                        continue

                    # Either a clean PNG, or an expected error PNG → use it
                    png_match = match
                    best_commands = [single_cmd]
                    break
                
                # If individual match found but no username tag, also try full sequence
                # This handles cases like DOCX=['display', 'username kacha1'] where full sequence 
                # may match a username-tagged PNG while individual matches clean PNG
                if png_match:
                    full_match = find_best_match(
                        node, match_commands, png_files, prefer_error=expect_error
                    )
                    if full_match and full_match != png_match:
                        # Full sequence matched a different (likely username-tagged) PNG
                        png_match = full_match
                        best_commands = match_commands
            else:
                # Normal block matching (full command sequence)
                png_match = find_best_match(
                    node, match_commands, png_files, prefer_error=expect_error
                )
                best_commands = match_commands

            if not png_match:
                results.append({
                    'node': node,
                    'block_idx': block_idx,
                    'para_idx': para_idx,
                    'commands': expanded_commands,
                    'match_path': None,
                    'action': 'no_match',
                })
                continue

            has_error = '[error]' in os.path.basename(png_match).lower()
            if has_error and not expect_error:
                # Error PNG found but block didn't expect it — skip
                results.append({
                    'node': node,
                    'block_idx': block_idx,
                    'para_idx': para_idx,
                    'commands': best_commands or expanded_commands,
                    'match_path': png_match,
                    'action': 'skipped_error',
                })
            else:
                # Either clean PNG, or error PNG on an expect_error block
                results.append({
                    'node': node,
                    'block_idx': block_idx,
                    'para_idx': para_idx,
                    'commands': best_commands or expanded_commands,
                    'match_path': png_match,
                    'action': 'inserted',
                })

    return results


def find_best_match(device: str, commands: list[str], png_files: list[str], prefer_error: bool = False) -> str | None:
    """Find the PNG whose filename matches the exact command sequence.

    The device name must match exactly (case-insensitive).
    Trailing 'quit' tokens are stripped from both cell commands and PNG filename
    before comparison, so cells with or without quit match correctly.
    The remaining command sequence must match exactly (same tokens, same order).
    
    Supports xxx.* placeholders in DOCX commands. For example, if the DOCX has
    'startup saved-configuration xxx.zip', it will match PNGs with any value
    where xxx.zip appears: 'startup saved-configuration backup.zip.png'.
    
    When prefer_error=True, prioritizes PNGs with [error] suffix over clean PNGs.
    """
    if not commands:
        return None

    # Regex to detect placeholder tokens like xxx.zip, xxx.cfg, etc.
    _PLACEHOLDER_RE = re.compile(r'^xxx\.[a-z]+$')

    # Build search tokens from device + commands
    search_tokens = sanitize_filename(device).lower().split()
    for cmd in commands:
        search_tokens.extend(sanitize_filename(cmd).lower().split())

    # Extract and remove "username VALUE" tag from search tokens (cell side)
    # The tag can appear anywhere in the command sequence, not just at the end.
    # We normalize by removing it entirely — it doesn't affect command matching.
    cell_username = None
    for i in range(len(search_tokens) - 1):
        if search_tokens[i] == 'username':
            cell_username = search_tokens[i + 1]
            # Remove both tokens
            del search_tokens[i:i + 2]
            break

    # Strip trailing 'quit' from search tokens (cell side)
    while len(search_tokens) > 1 and search_tokens[-1] == 'quit':
        search_tokens.pop()

    cmd_tokens = search_tokens[1:]  # Skip device name
    if not cmd_tokens:
        # Device-only match
        for png_path in png_files:
            png_name = os.path.basename(png_path).replace('.png', '').lower()
            png_tokens = png_name.split()
            if png_tokens and png_tokens[0] == search_tokens[0]:
                return png_path
        return None

    # Canonicalize DOCX command tokens via abbreviation expansion
    # (same as PNG side, so 'dis cur' ↔ 'display current-configuration' matches)
    cmd_tokens = expand_abbreviations([' '.join(cmd_tokens)])[0].split()

    best_match = None
    current_best_score = None

    for png_path in png_files:
        png_name = os.path.basename(png_path).replace('.png', '').lower()
        png_tokens = png_name.split()

        # Device name must match exactly as first token
        if not png_tokens or png_tokens[0] != search_tokens[0]:
            continue

        # Strip trailing 'quit' from PNG tokens (PNG side)
        png_cmd_tokens = png_tokens[1:]
        while png_cmd_tokens and png_cmd_tokens[-1] in ('quit', '[error]'):
            png_cmd_tokens.pop()

        # Extract and remove "username VALUE" tag from PNG tokens
        # The tag can appear anywhere in the filename, not just at the end.
        png_username = None
        for i in range(len(png_cmd_tokens) - 1):
            if png_cmd_tokens[i] == 'username':
                png_username = png_cmd_tokens[i + 1]
                # Remove both tokens
                del png_cmd_tokens[i:i + 2]
                break

        # Username filter: DOCX and PNG must have the same username state
        # - Both have no username → OK
        # - Both have same username → OK
        # - Mismatch (one has, one doesn't, or different values) → skip
        if cell_username != png_username:
            continue

        # Canonicalize PNG command tokens via same abbreviation expansion as DOCX
        if png_cmd_tokens:
            png_cmd_str = ' '.join(png_cmd_tokens)
            png_cmd_tokens = expand_abbreviations([png_cmd_str])[0].split()

        # Check token count matches
        if len(png_cmd_tokens) != len(cmd_tokens):
            continue

        # Compare token-by-token, with xxx.* placeholder support
        all_match = True
        for pt, ct in zip(png_cmd_tokens, cmd_tokens):
            placeholder_match = _PLACEHOLDER_RE.match(ct)
            if placeholder_match:
                # DOCX has xxx.* placeholder — compare the file extension only
                # e.g., xxx.zip in DOCX should match xxx.zip in PNG, but NOT xxx.cfg
                docx_ext = ct.split('.')[-1]
                if '.' not in pt:
                    all_match = False
                    break
                png_ext = pt.split('.')[-1]
                if png_ext != docx_ext:
                    all_match = False
                    break
                continue
            if pt != ct:
                all_match = False
                break

        if not all_match:
            continue

        # Scoring: when prefer_error, error PNGs win first
        # Username matching is handled by the continue check above;
        # by this point usernames are guaranteed to match.
        is_error = '[error]' in png_name

        if prefer_error:
            score = (
                0 if is_error else 1,
                len(png_tokens)
            )
        else:
            score = (
                0 if not is_error else 1,
                len(png_tokens)
            )
        if best_match is None or score < current_best_score:
            best_match = png_path
            current_best_score = score

    return best_match


# --- Main ---
if __name__ == "__main__":
    import sys

    cfg = load_config()
    docx_input = cfg.get("docx_input", _DEFAULT_CONFIG["docx_input"])
    png_path = cfg.get("png_path", _DEFAULT_CONFIG["png_path"])
    docx_output = cfg.get("docx_output", _DEFAULT_CONFIG["docx_output"])
    target_sections = cfg.get("target_sections", _DEFAULT_CONFIG["target_sections"])

    if not os.path.exists(docx_input):
        print(f"❌ Error: DOCX input file not found: {docx_input}")
        sys.exit(1)

    try:
        document = Document(docx_input)
    except Exception as e:
        print(f"❌ Error: Failed to open DOCX file: {e}")
        sys.exit(1)

    png_files = sorted(glob.glob(png_path))
    if not png_files:
        print(f"⚠ Warning: No PNG files found at: {png_path}")
        print("  Run process_network_logs.py first to generate screenshots.")

    # Collect marker/txt pairs during the main loop for OLE embedding
    marker_txt_pairs = []

    print(f"Loaded: {docx_input}")
    print(f"Found {len(png_files)} PNG files in {png_path}")

    # Permit logic — controls which test cases get images
    # (customizable per document)
    PERMIT_FALSE_KEYWORDS = (
        'Reboot device.',
        'To verify power module resetting on a switch',
        'Reinsert the power module into slot 1-',
        'To verify fan module resetting on a CE',
    )
    PERMIT_TRUE_KEYWORDS = (
        'T01-02',
        'Simulate alarms and view alarms on NMS',
        'T01-05',
    )

    for table_index, table in enumerate(document.tables):
        if target_sections:
            if not should_process_table(document, table_index, target_sections):
                continue

        print(f"\n--- Table {table_index + 1} ---")

        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text

                # Determine permit
                permit = True
                for kw in PERMIT_FALSE_KEYWORDS:
                    if kw in cell_text:
                        permit = False
                for kw in PERMIT_TRUE_KEYWORDS:
                    if kw in cell_text:
                        permit = True

                if not permit:
                    continue

                # Option B: Block-scoped node matching
                match_results = match_cell_blocks(cell.paragraphs, png_files)

                if not match_results:
                    continue

                # Track inserted paragraphs to prevent double-insertion
                inserted_paragraphs = set()
                # Track paragraph index shift due to OLE marker insertions
                para_idx_shift = 0

                for result in match_results:
                    if result['action'] != 'inserted':
                        if result['action'] == 'skipped_error':
                            print(f"  Skipped (error): {result['node']} block {result['block_idx']}")
                        else:
                            print(f"  No match: {result['node']} block {result['block_idx']}")
                        continue

                    para_idx = result['para_idx'] + para_idx_shift
                    if para_idx in inserted_paragraphs:
                        continue  # Safety: don't double-insert

                    paragraph = cell.paragraphs[para_idx]
                    match_path = result['match_path']
                    commands = result['commands']

                    paragraph.paragraph_format.first_line_indent = 0
                    paragraph.paragraph_format.left_indent = 0
                    run = paragraph.add_run()
                    run.add_break()
                    run.add_break()
                    run = paragraph.add_run()
                    run.add_picture(match_path, width=Inches(6.495))
                    print(f"  Inserted: {os.path.basename(match_path)}")
                    print(f"    Node: {result['node']}, Block: {result['block_idx']}")
                    print(f"    Command: {' '.join(commands)}")
                    inserted_paragraphs.add(para_idx)

                    # Check if there's a full-output .txt for this truncated command
                    txt_path = match_path.replace('.png', '.log')
                    if os.path.exists(txt_path):
                        marker = f"OLE_MARKER_{uuid.uuid4().hex[:12]}"
                        # Insert marker paragraph immediately AFTER the image paragraph
                        # using XML addnext so it appears right below the image
                        from docx.oxml import OxmlElement
                        new_p = OxmlElement('w:p')
                        new_r = OxmlElement('w:r')
                        new_t = OxmlElement('w:t')
                        new_t.text = marker
                        new_r.append(new_t)
                        new_p.append(new_r)
                        paragraph._element.addnext(new_p)
                        print(f"    Marked for OLE embedding: {os.path.basename(txt_path)}")
                        marker_txt_pairs.append((marker, txt_path))
                        # Increment shift counter - all subsequent para_idx need +1
                        para_idx_shift += 1

    if marker_txt_pairs:
        # Save to temp first
        docx_temp = docx_output + ".tmp"
        document.save(docx_temp)
        print(f"\nSaved temp: {docx_temp}")

        # Post-process: embed .txt attachments as OLE objects via Word COM
        import tempfile
        tmp_dir = tempfile.mkdtemp(prefix="putpnginword_")
        tmp_final = os.path.join(tmp_dir, os.path.basename(docx_output))
        shutil.copy2(docx_temp, tmp_final)
        try:
            embed_txt_via_word(docx_temp, marker_txt_pairs)
            shutil.move(docx_temp, docx_output)
            print(f"\nSaved: {docx_output}")
        except ImportError as e:
            print(f"\n⚠ Warning: OLE embedding skipped — {e}")
            print("  Either pywin32 is not installed or Microsoft Word is not available.")
            print(f"  Your document has been saved without embedded .txt files: {docx_output}")
            shutil.copy2(docx_temp, docx_output)
        except Exception as e:
            print(f"\n⚠ Warning: OLE embedding failed — {e}")
            print(f"  Your document has been saved without embedded .txt files: {docx_output}")
            shutil.copy2(docx_temp, docx_output)
        finally:
            # Clean up temp files
            if os.path.exists(docx_temp):
                os.remove(docx_temp)
            if os.path.exists(tmp_final):
                os.remove(tmp_final)
            if os.path.exists(tmp_dir):
                os.rmdir(tmp_dir)
    else:
        document.save(docx_output)
        print(f"\nSaved: {docx_output}")
