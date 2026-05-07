"""Diagnose issues in comprehensive_test_cases_updated_RESULT.docx"""
from docx import Document
from putpnginword import parse_paragraphs_detailed, _merge_empty_blocks

doc = Document('comprehensive_test_cases_updated_RESULT.docx')

print("=" * 60)
print("TABLE 13: Error + Username Case")
print("=" * 60)
table13 = doc.tables[12]  # Table 13 (0-indexed)
cell = table13.rows[1].cells[1]

print("\nCell paragraphs:")
for pi, p in enumerate(cell.paragraphs):
    if p.text.strip():
        print(f"  [{pi}]: {p.text.strip()}")

print("\nParsed blocks:")
blocks = parse_paragraphs_detailed(cell.paragraphs)
for i, (cmds, nodes, err) in enumerate(blocks):
    print(f"  Block {i}: cmds={cmds}, nodes={[n[0] for n in nodes]}, error={err}")

merged = _merge_empty_blocks(blocks)
print("\nMerged blocks:")
for i, (cmds, nodes, err) in enumerate(merged):
    print(f"  Block {i}: cmds={cmds}, nodes={[n[0] for n in nodes]}, error={err}")

print("\n" + "=" * 60)
print("TABLE 15: Image Positioning")
print("=" * 60)
table15 = doc.tables[14]
cell15 = table15.rows[1].cells[1]

print("\nCell paragraphs (showing runs/images):")
for pi, p in enumerate(cell15.paragraphs):
    text = p.text.strip()
    has_image = len(p._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) > 0
    print(f"  [{pi}]: text='{text[:60]}' image={has_image}")

print("\n" + "=" * 60)
print("TABLE 19: Image Positioning")
print("=" * 60)
table19 = doc.tables[18]
cell19 = table19.rows[1].cells[1]

print("\nCell paragraphs:")
for pi, p in enumerate(cell19.paragraphs):
    text = p.text.strip()
    has_image = len(p._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')) > 0
    print(f"  [{pi}]: text='{text[:60]}' image={has_image}")

print("\n" + "=" * 60)
print("TXT ATTACHMENTS (OLE)")
print("=" * 60)
# Count OLE objects
count = 0
for rel in doc.part.rels.values():
    if 'oleObject' in rel.target_ref:
        count += 1
print(f"Total OLE objects in document: {count}")

# Check which tables have images
for ti in [14, 18]:
    table = doc.tables[ti]
    cell = table.rows[1].cells[1]
    img_count = 0
    for p in cell.paragraphs:
        if p._p.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
            img_count += 1
    print(f"Table {ti+1}: {img_count} images in cell")
