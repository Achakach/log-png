"""Add 'Expected Result' column to comprehensive_test_cases.docx"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

doc = Document('comprehensive_test_cases.docx')

# Expected results for each of the 19 tables
expected_results = [
    # Table 1 - Basic single command
    "TUC-TEST01 display device.png\nTUC-TEST02 display device.png\nTUC-TEST03 display device.png",
    # Table 2 - Nested block
    "TUC-TEST01 system-view interface GigabitEthernet0_0_1 display this quit quit.png",
    # Table 3 - Multiple commands
    "TUC-TEST01/02/03 display device.png\nTUC-TEST01/02/03 display clock.png",
    # Table 4 - Different nodes
    "TUC-TEST02 display device.png\nTUC-TEST01 display clock.png",
    # Table 5 - No match
    "No match (display cpu-usage not in logs)",
    # Table 6 - Multiple independent
    "TUC-TEST01 display device.png\nTUC-TEST01 display clock.png",
    # Table 7 - Nested without full quit
    "TUC-TEST01 system-view interface GigabitEthernet0_0_1 display this quit quit.png",
    # Table 8 - Empty cell
    "No match (no nodes present)",
    # Table 9 - Wrong command
    "Skipped (error PNG but no error text)",
    # Table 10 - Mixed nested + standalone
    "TUC-TEST01/02/03 system-view interface GigabitEthernet0_0_1 display this quit quit.png\nTUC-TEST01/02/03 display device.png\nTUC-TEST01/02/03 display clock.png",
    # Table 11 - Skip error match next
    "Skipped (wrong-command)\nTUC-TEST01/02/03 display device.png",
    # Table 12 - Model labels
    "TUC-TEST01/02/03 display device.png\nTUC-TEST01/02/03 display clock.png",
    # Table 13 - Error with typo
    "TUC-TEST01/02/03 display current-configutation [error].png",
    # Table 14 - Placeholder
    "TUC-TEST01/02/03 startup saved-configuration backup.zip.png\nTUC-TEST01/02/03 startup saved-configuration test.cfg.png\nTUC-TEST01/02/03 startup saved-configuration file.hi.png",
    # Table 15 - Clean display current-configuration
    "TUC-TEST01/02/03 display current-configuration.png",
    # Table 16 - Error system-view
    "TUC-TEST01/02/03 system-view [error].png",
    # Table 17 - Error display current-configuration
    "TUC-TEST01/02/03 display current-configuration [error].png",
    # Table 18 - Username kacha1
    "TUC-TEST01/02/03 display current-configuration username kacha1.png",
    # Table 19 - Clean display current-configuration
    "TUC-TEST01/02/03 display current-configuration.png",
]


def insert_cell_after(row, source_idx=0):
    """Insert a new cell after the source_idx cell in a row."""
    source_cell = row.cells[source_idx]
    new_tc = deepcopy(source_cell._tc)
    # Reset text content
    for p in new_tc.findall('.//' + qn('w:p')):
        p.clear()
        p.append(OxmlElement('w:pPr'))
    source_cell._tc.addnext(new_tc)


# Process each table
for ti, table in enumerate(doc.tables):
    # Add header cell (Expected Result)
    row0 = table.rows[0]
    
    # Check if table already has 3 columns
    if len(row0.cells) >= 3:
        print(f"Table {ti+1}: Already has 3 columns, skipping")
        continue
    
    # Insert new column (after Cell Content column, index 1)
    for row in table.rows:
        source_cell = row.cells[-1]
        new_tc = deepcopy(source_cell._tc)
        # Clear text
        for p in new_tc.findall('.//' + qn('w:p')):
            # Clear runs but keep paragraph
            for r in list(p.findall('.//' + qn('w:r'))):
                p.remove(r)
        source_cell._tc.addnext(new_tc)
    
    # Set header text
    header_row = table.rows[0]
    new_header_cell = header_row.cells[-1]  # Last cell is new
    new_header_cell.paragraphs[0].clear()
    new_header_cell.paragraphs[0].add_run("Expected Result")
    
    # Set expected result text
    data_row = table.rows[1]
    new_data_cell = data_row.cells[-1]  # Last cell is new
    new_data_cell.paragraphs[0].clear()
    
    # Add expected result text with line breaks
    lines = expected_results[ti].split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            new_data_cell.add_paragraph(line)
        else:
            new_data_cell.paragraphs[0].add_run(line)
    
    print(f"Table {ti+1}: Added 'Expected Result' column")

# Save with new name
doc.save('comprehensive_test_cases_updated.docx')
print("\nSaved: comprehensive_test_cases_updated.docx")
print(f"Tables processed: {len(doc.tables)}")
