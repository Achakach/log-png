from docx import Document
from putpnginword import parse_paragraphs_detailed, _merge_empty_blocks, find_best_match, expand_abbreviations
import glob

# Create a DOCX with ALL 19 cases in ONE cell
doc = Document()
table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]

# Add all 19 cases as paragraphs in one cell
cell.add_paragraph("=== All 19 Cases in One Cell ===")
cell.add_paragraph("")

# Case 1: Basic
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 2: Nested
cell.add_paragraph("<HUAWEI>system-view")
cell.add_paragraph("[HW]interface GigabitEthernet0/0/1")
cell.add_paragraph("[HW-GigabitEthernet0/0/1]display this")
cell.add_paragraph("[HW-GigabitEthernet0/0/1]quit")
cell.add_paragraph("[HW]quit")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("")

# Case 3: Multiple
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("<HUAWEI>display clock")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 4: Different nodes
cell.add_paragraph("<TUC-TEST01>display device")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST01>display clock")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("")

# Case 5: Missing
cell.add_paragraph("<TUC-TEST01>display cpu-usage")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("")

# Case 6: Multiple independent
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<HUAWEI>display clock")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("")

# Case 7: Nested without quit
cell.add_paragraph("<HUAWEI>system-view")
cell.add_paragraph("[HW]interface GigabitEthernet0/0/1")
cell.add_paragraph("[HW-GigabitEthernet0/0/1]display this")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("")

# Case 8: Empty
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("")

# Case 9: Wrong command
cell.add_paragraph("<HUAWEI>display wrong-command")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("")

# Case 10: Mixed
cell.add_paragraph("<HUAWEI>system-view")
cell.add_paragraph("[HW]interface GigabitEthernet0/0/1")
cell.add_paragraph("[HW-GigabitEthernet0/0/1]display this")
cell.add_paragraph("[HW-GigabitEthernet0/0/1]quit")
cell.add_paragraph("[HW]quit")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("<HUAWEI>display clock")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 11: Skip error
cell.add_paragraph("<HUAWEI>display wrong-command")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 12: Model labels
cell.add_paragraph("CEx01")
cell.add_paragraph("<HUAWEI>display device")
cell.add_paragraph("CEx02")
cell.add_paragraph("<HUAWEI>display clock")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 13: Error with typo
cell.add_paragraph("<HUAWEI>display current-configutation")
cell.add_paragraph("Error: Do not have permission")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 14: Placeholder
cell.add_paragraph("<HUAWEI>startup saved-configuration xxx.zip")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("<HUAWEI>startup saved-configuration xxx.cfg")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("<HUAWEI>startup saved-configuration xxx.hi")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 15: Clean
cell.add_paragraph("<HUAWEI>display current-configuration")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 16: Error system-view
cell.add_paragraph("<HUAWEI>system-view")
cell.add_paragraph("Error: you have no right")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 17: Error display current-configuration
cell.add_paragraph("<HUAWEI>display current-configuration")
cell.add_paragraph("Error: you have no right")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 18: Username
cell.add_paragraph("<HUAWEI>display current-configuration")
cell.add_paragraph("<HUAWEI>username kacha1")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")
cell.add_paragraph("")

# Case 19: Clean again
cell.add_paragraph("<HUAWEI>display current-configuration")
cell.add_paragraph("<TUC-TEST01>")
cell.add_paragraph("<TUC-TEST02>")
cell.add_paragraph("<TUC-TEST03>")

doc.save("all_in_one_cell.docx")
print("Created: all_in_one_cell.docx")
