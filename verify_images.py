from docx import Document
import sys
sys.path.insert(0, '.')
from putpnginword import parse_paragraphs

result_doc = Document('testthisplease_RESULT_v4.docx')
orig_doc = Document('testthisplease.docx')

print('=== IMAGE INSPECTION REPORT ===\n')

for ti, (table, orig_table) in enumerate(zip(result_doc.tables, orig_doc.tables)):
    print(f'Table {ti+1}:')
    for ri, (row, orig_row) in enumerate(zip(table.rows, orig_table.rows)):
        for ci, (cell, orig_cell) in enumerate(zip(row.cells, orig_row.cells)):
            text = orig_cell.text.strip()
            if not text or len(text) <= 10:
                continue
                
            img_count = len(cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
            
            blocks = parse_paragraphs(orig_cell.paragraphs)
            total_nodes = sum(len(nodes) for _, nodes in blocks)
            
            if total_nodes > 0:
                status = 'OK' if img_count >= total_nodes else f'MISSING {total_nodes - img_count}'
                print(f'  Row {ri} Col {ci}: {img_count} image(s) / {total_nodes} nodes -> {status}')
                # Show blocks
                for bi, (cmds, nodes) in enumerate(blocks):
                    print(f'    Block {bi}: {len(nodes)} nodes: {nodes}')
                
                # Show text preview
                preview = text[:120].replace('\n', ' | ')
                print(f'    Text: {preview}')
                print()

print('=== END REPORT ===')
