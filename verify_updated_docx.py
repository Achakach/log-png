from docx import Document

doc = Document('comprehensive_test_cases_updated_RESULT.docx')
print(f"Total tables: {len(doc.tables)}")

for ti, table in enumerate(doc.tables):
    print(f"\n=== Table {ti+1} ===")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.paragraphs[0].text.strip() if cell.paragraphs else ''
            if text:
                print(f"  [{ri}][{ci}]: {text[:100]}")
