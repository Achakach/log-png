"""Verify result DOCX image dimensions"""
from docx import Document
from PIL import Image
import os

doc = Document("comprehensive_test_cases_RESULT_v6.docx")

print("=== RESULT DOCX IMAGE VERIFICATION ===")
print(f"Total tables: {len(doc.tables)}")
print()

for table_idx, table in enumerate(doc.tables):
    print(f"--- Table {table_idx + 1} ---")
    for row_idx, row in enumerate(table.rows):
        for cell in enumerate(row.cells):
            cell = row.cells[1]  # Cell Content column
            
            if cell.text.strip():
                # Count images in this cell
                image_count = 0
                image_heights = []
                
                for para in cell.paragraphs:
                    for run in para.runs:
                        # Check for inline shapes (images)
                        if run._element.xpath('.//a:blip'):
                            image_count += 1
                            # Try to get image dimensions from the relationship
                            try:
                                blips = run._element.xpath('.//a:blip')
                                for blip in blips:
                                    embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                    if embed:
                                        image_part = doc.part.related_parts[embed]
                                        image_bytes = image_part.blob
                                        # Save temp and check size
                                        temp_path = f"temp_img_{table_idx}_{image_count}.png"
                                        with open(temp_path, 'wb') as f:
                                            f.write(image_bytes)
                                        img = Image.open(temp_path)
                                        image_heights.append(f"{img.size[0]}x{img.size[1]} ({img.size[1]/96:.1f}in)")
                                        os.remove(temp_path)
                            except Exception as e:
                                image_heights.append("unknown")
                
                if image_count > 0:
                    print(f"  Images: {image_count}")
                    for i, h in enumerate(image_heights):
                        print(f"    Image {i+1}: {h}")
                break
    print()

print("Done! Open comprehensive_test_cases_RESULT_v6.docx in Word to verify.")
