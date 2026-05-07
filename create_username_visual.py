from docx import Document

doc = Document()

doc.add_heading('Username Tag Test - Both Scenarios', level=1)

# Scenario A: DOCX no username, PNG has username metadata tag
doc.add_paragraph("Scenario A: DOCX has no username, PNG has username metadata (stripped during match)")
table_a = doc.add_table(rows=1, cols=1)
cell_a = table_a.rows[0].cells[0]
cell_a.add_paragraph("<TUCxx01>dis cur")
cell_a.add_paragraph("<TUCxx01>")

# Scenario B: DOCX has username, PNG has username as real command
doc.add_paragraph("Scenario B: DOCX has username + command, PNG matches exactly")
table_b = doc.add_table(rows=1, cols=1)
cell_b = table_b.rows[0].cells[0]
cell_b.add_paragraph("<TUCxx01>username kacha1")
cell_b.add_paragraph("<TUCxx01>display current-configuration")
cell_b.add_paragraph("<TUCxx01>")

# Scenario C: Error case with username
doc.add_paragraph("Scenario C: Error case with username")
table_c = doc.add_table(rows=1, cols=1)
cell_c = table_c.rows[0].cells[0]
cell_c.add_paragraph("<TUCxx01>dis cur")
cell_c.add_paragraph("Error: Unrecognized command")
cell_c.add_paragraph("<TUCxx01>")

doc.save("username_visual_test.docx")
print("Created username_visual_test.docx")
