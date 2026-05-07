"""Convert example.txt into a proper DOCX for testing."""
from docx import Document
from docx.shared import Pt

doc = Document()

with open('example.txt', 'r', encoding='utf-8') as f:
    lines = f.read().splitlines()

i = 0
while i < len(lines):
    line = lines[i].strip()

    # Skip empty lines
    if not line:
        i += 1
        continue

    # Skip section headers
    if line in ['dis cur example!']:
        doc.add_heading(line, level=2)
        i += 1
        continue

    # Check if this is a DOCX cell block
    # Pattern: <HUAWEI>... followed by <TUCxx> nodes
    if line.startswith('<HUAWEI>') or line.startswith('<TUCxx>'):
        # Collect block lines until next empty line or new section
        block_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                i += 1
                continue
            if l == 'dis cur example!':
                break
            block_lines.append(l)
            i += 1

        # Create table for this block
        if block_lines:
            table = doc.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            for bl in block_lines:
                cell.add_paragraph(bl)
            # Add an extra empty paragraph for spacing
            doc.add_paragraph()
    else:
        i += 1

doc.save('example_as_docx.docx')
print("Created: example_as_docx.docx")
