from docx import Document
import sys, os
sys.path.insert(0, '.')
from putpnginword import parse_paragraphs_detailed, _merge_empty_blocks, expand_abbreviations, find_best_match

result_doc = Document('comprehensive_test_cases_RESULT.docx')

print('=== DETAILED MATCHING TRACE ===\n')

for ti, table in enumerate(result_doc.tables):
    case_num = ti + 1
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if ci == 1 and ri == 1:  # Content row
                text = cell.text.strip()
                if not text or len(text) <= 5:
                    continue
                
                # Count images
                img_count = len(cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
                
                # Get PNG files
                png_files = [os.path.join('screenshots', f) for f in os.listdir('screenshots') if f.endswith('.png')]
                
                # Parse blocks
                blocks = parse_paragraphs_detailed(cell.paragraphs)
                merged_blocks = _merge_empty_blocks(blocks)
                was_merged = len(merged_blocks) != len(blocks)
                
                print(f'--- Case {case_num} ---')
                print(f'Images in cell: {img_count}')
                
                if was_merged:
                    print('  MERGE APPLIED: Empty blocks merged into next block')
                
                print(f'  Parsed blocks ({len(blocks)}):')
                for bi, (cmds, nodes) in enumerate(blocks):
                    node_list = [n for n,_ in nodes]
                    print(f'    Block {bi}: {cmds} -> nodes={node_list}')
                
                if was_merged:
                    print(f'  After merge ({len(merged_blocks)}):')
                    for bi, (cmds, nodes) in enumerate(merged_blocks):
                        node_list = [n for n,_ in nodes]
                        print(f'    Block {bi}: {cmds} -> nodes={node_list}')
                
                # Trace matching for each node
                for bi, (commands, block_nodes) in enumerate(merged_blocks):
                    if not commands or not block_nodes:
                        continue
                    
                    expanded = expand_abbreviations(commands)
                    
                    for node, para_idx in block_nodes:
                        print(f'\n    {node} (para {para_idx}):')
                        
                        if was_merged:
                            # Try each command individually
                            for cmd in expanded:
                                match = find_best_match(node, [cmd], png_files)
                                if match:
                                    has_error = '[error]' in os.path.basename(match).lower()
                                    status = '❌ [ERROR] SKIP' if has_error else '✅ INSERT'
                                    print(f'      Try "{cmd}" -> {os.path.basename(match)} {status}')
                                    if not has_error:
                                        break
                                else:
                                    print(f'      Try "{cmd}" -> NO MATCH')
                        else:
                            # Normal full command matching
                            match = find_best_match(node, expanded, png_files)
                            if match:
                                has_error = '[error]' in os.path.basename(match).lower()
                                status = '❌ [ERROR] SKIP' if has_error else '✅ INSERT'
                                print(f'      Try {expanded} -> {os.path.basename(match)} {status}')
                            else:
                                print(f'      Try {expanded} -> NO MATCH')
                
                print()
