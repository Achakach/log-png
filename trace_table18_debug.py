import glob
from putpnginword import parse_paragraphs_detailed, _merge_empty_blocks, expand_abbreviations, find_best_match
from docx import Document

doc = Document('comprehensive_test_cases.docx')
table = doc.tables[17]
cell = table.rows[1].cells[1]

print("=== Cell paragraphs ===")
for pi, p in enumerate(cell.paragraphs):
    if p.text.strip():
        print(f"  [{pi}]: {p.text.strip()}")

blocks = parse_paragraphs_detailed(cell.paragraphs)
print(f"\nBlocks before merge: {len(blocks)}")
for i, (cmds, nodes, err) in enumerate(blocks):
    print(f"  Block {i}: cmds={cmds}, nodes={nodes}, error={err}")

merged = _merge_empty_blocks(blocks)
print(f"\nBlocks after merge: {len(merged)}")
for i, (cmds, nodes, err) in enumerate(merged):
    print(f"  Block {i}: cmds={cmds}, nodes={nodes}, error={err}")

png_files = glob.glob('screenshots/*.png')

# Simulate matching
for block_idx, (commands, block_nodes, expect_error) in enumerate(merged):
    print(f"\n=== Matching Block {block_idx} ===")
    expanded_cmds = expand_abbreviations(commands)
    print(f"Commands: {expanded_cmds}")
    
    for node, para_idx in block_nodes:
        png_match = find_best_match(node, expanded_cmds, png_files)
        if png_match:
            print(f"  Node {node}: MATCH {png_match.split(chr(92))[-1]}")
        else:
            print(f"  Node {node}: NO MATCH")
