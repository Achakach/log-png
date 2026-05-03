from docx import Document
import sys
sys.path.insert(0, '.')
from putpnginword import parse_paragraphs_detailed

doc = Document('document structure example.docx')

print(f"Document has {len(doc.tables)} tables\n")

for ti, table in enumerate(doc.tables):
    print(f"=== Table {ti+1} ===")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text and len(text) > 5:
                blocks = parse_paragraphs_detailed(cell.paragraphs)
                if blocks:
                    print(f"  Row {ri} Col {ci}: {len(blocks)} block(s)")
                    for bi, (cmds, nodes) in enumerate(blocks):
                        print(f"    Block {bi}: cmds={cmds}, nodes={nodes}")
                    preview = text[:200].replace('\n', ' | ')
                    print(f"    Text: {preview}\n")
