from docx import Document

doc = Document()

doc.add_heading('SSH Error Test', level=1)

# Standard case: command failed, no error flag in word
table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]
cell.add_paragraph("<TUCxx01>dis cur")
cell.add_paragraph("<TUCxx01>")

doc.save("ssh_error_test.docx")
print("Created ssh_error_test.docx")
