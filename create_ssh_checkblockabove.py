from docx import Document

doc = Document()

# Heading
doc.add_heading('SSH Merge Test from checkblockabove.txt', level=1)

# Replicate exactly what the log file produces
# The DOCX should ONLY have the display command (no stelnet in cell)
doc.add_paragraph("This test matches: screenshots\\TUCxx01 dis cur.png")

table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]
cell.add_paragraph("<TUCxx01>dis cur")
cell.add_paragraph("<TUCxx01>")

doc.save("ssh_merge_checkblockabove.docx")
print("Created ssh_merge_checkblockabove.docx")
