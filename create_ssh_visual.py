from docx import Document
from docx.shared import Pt

doc = Document()

# Heading
doc.add_heading('SSH Merge Visual Test', level=1)

# Case 1: SSH + dis cur
doc.add_heading('Case 1: stelnet + dis cur', level=2)
table1 = doc.add_table(rows=1, cols=1)
cell1 = table1.rows[0].cells[0]
cell1.add_paragraph("<OSTORxx022>stelnet x.x.x.x 22")
cell1.add_paragraph("<TUCxx01>dis cur")
cell1.add_paragraph("<TUCxx01>")

# Case 2: SSH + display current-configuration
doc.add_heading('Case 2: ssh + display current-configuration', level=2)
table2 = doc.add_table(rows=1, cols=1)
cell2 = table2.rows[0].cells[0]
cell2.add_paragraph("<OSTORxx022>ssh admin@10.1.1.2")
cell2.add_paragraph("<TUCxx02>display current-configuration")
cell2.add_paragraph("<TUCxx02>")

doc.save("ssh_merge_visual_test.docx")
print("Created ssh_merge_visual_test.docx")
