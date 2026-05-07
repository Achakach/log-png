import glob
from putpnginword import match_cell_blocks
from docx import Document

doc = Document('comprehensive_test_cases.docx')
table = doc.tables[17]  # Table 18
cell = table.rows[1].cells[1]
paras = cell.paragraphs

png_files = glob.glob('screenshots/*.png')
results = match_cell_blocks(paras, png_files)

print('=== Table 18 Results ===')
for r in results:
    node = r['node']
    action = r['action']
    match = r['match_path']
    cmds = r['commands']
    
    if match:
        png_name = match.split(chr(92))[-1]
        print(f'  {action.upper()}: {png_name}')
        print(f'    Node: {node}, Commands: {cmds}')
    else:
        print(f'  NO MATCH: {node}')
        print(f'    Commands: {cmds}')
