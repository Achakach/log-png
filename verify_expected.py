"""Verify putpnginword.py output matches Expected Result column.

Reads comprehensive_test_cases_updated.docx, compares Expected Result
to what was actually generated in comprehensive_test_cases_updated_RESULT_v2.docx.
"""
from docx import Document
import os

SRC_DOCX = 'comprehensive_test_cases_updated.docx'
RESULT_DOCX = 'comprehensive_test_cases_updated_RESULT_v2.docx'


def read_table_data(doc_path):
    """Read all tables: list of {case, cell_text, expected}."""
    doc = Document(doc_path)
    tables = []
    for ti, table in enumerate(doc.tables):
        rows = []
        for ri, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if len(cells) >= 3:
                rows.append({
                    'table': ti + 1,
                    'row': ri + 1,
                    'case': cells[0],
                    'content': cells[1],
                    'expected': cells[2],
                })
        tables.append(rows)
    return tables


def count_images_in_cell(cell):
    """Count images embedded in a cell's paragraphs."""
    count = 0
    for para in cell.paragraphs:
        for run in para.runs:
            if run._element.xpath('.//a:blip'):
                count += 1
            # Also check for inline shapes
            count += len(run._element.xpath('.//wp:inline'))
    return count


def count_images_in_docx(doc_path):
    """Count total images in docx."""
    doc = Document(doc_path)
    total = 0
    for rel in doc.part.rels.values():
        if 'image' in rel.reltype:
            total += 1
    return total


def main():
    print(f"Source: {SRC_DOCX}")
    print(f"Result: {RESULT_DOCX}")
    print()

    if not os.path.exists(RESULT_DOCX):
        print("❌ Result docx not found. Run putpnginword.py first.")
        return

    src_tables = read_table_data(SRC_DOCX)
    result_doc = Document(RESULT_DOCX)

    # Compare each table
    total_expected = 0
    total_matched = 0
    total_mismatch = 0

    for ti, src_rows in enumerate(src_tables):
        if ti >= len(result_doc.tables):
            print(f"Table {ti+1}: Missing in result docx")
            continue

        result_table = result_doc.tables[ti]
        print(f"--- Table {ti+1} ---")

        for ri, src_row in enumerate(src_rows):
            if ri >= len(result_table.rows):
                print(f"  Row {ri+1}: Missing in result")
                continue

            expected = src_row['expected']
            content = src_row['content']
            case = src_row['case']

            # Count images in result cell (col 2 = content cell)
            result_cell = result_table.rows[ri].cells[1] if len(result_table.rows[ri].cells) > 1 else None
            img_count = count_images_in_cell(result_cell) if result_cell else 0

            # Parse expected
            expected_lines = [l.strip() for l in expected.split('\n') if l.strip()]
            is_error = any('error' in l.lower() for l in expected_lines)
            is_skip = any('skip' in l.lower() or 'no match' in l.lower() for l in expected_lines)
            expected_count = sum(1 for l in expected_lines if l.endswith('.png'))

            # Determine status
            if is_skip:
                if img_count == 0:
                    status = "✅"
                    total_matched += 1
                else:
                    status = "❌"
                    total_mismatch += 1
            elif expected_count > 0:
                if img_count == expected_count:
                    status = "✅"
                    total_matched += 1
                else:
                    status = f"⚠️  Expected {expected_count}, got {img_count}"
                    total_mismatch += 1
            else:
                status = "➖"

            total_expected += 1

            if status not in ("✅", "➖"):
                print(f"  {status} {case}: expected={expected_count}, got={img_count}")
                print(f"     Content: {content[:60]}...")
                print(f"     Expected: {expected[:80]}...")
            elif status == "✅":
                print(f"  ✅ {case}: {img_count} image(s) inserted")

    print()
    print(f"Summary: {total_matched}/{total_expected} matched, {total_mismatch} mismatch")


if __name__ == '__main__':
    main()
