import sys
sys.path.insert(0, '.')
from putpnginword import parse_paragraphs, match_cell_blocks, expand_abbreviations, find_best_match
from docx import Document
import re, os

png_files = [os.path.join('screenshots', f) for f in os.listdir('screenshots') if f.endswith('.png')]

doc = Document('testthisplease.docx')
print(f'Total PNG files: {len(png_files)}')
print()

for ti, table in enumerate(doc.tables):
    print(f'=== Table {ti+1} ===')
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text and len(text) > 10:
                blocks = parse_paragraphs(cell.paragraphs)
                if not blocks:
                    continue
                    
                img_count = len(cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
                
                # Check which blocks match
                print(f'\nRow {ri} Col {ci}:')
                print(f'  Cell text preview: {text[:150].replace("\n", " | ")}')
                print(f'  Currently has {img_count} image(s)')
                
                for bi, (cmds, nodes) in enumerate(blocks):
                    expanded = expand_abbreviations(cmds)
                    print(f'    Block {bi}: cmds={expanded}, nodes={nodes}')
                    for node in nodes:
                        match = find_best_match(node, expanded, png_files)
                        if match:
                            has_error = '[error]' in os.path.basename(match).lower()
                            status = 'OK' if not has_error else 'ERROR (skip!)'
                            print(f'      {node} -> {os.path.basename(match)} [{status}]')
                        else:
                            print(f'      {node} -> NO MATCH')
