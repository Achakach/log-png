"""Generate a test Word document for testing putpnginword.py.

Creates a .docx with tables containing single and nested command cells
that match the format expected by putpnginword.py.

Each line in a cell is a separate paragraph (not \n within one string)
so that putpnginword.py can parse them correctly via cell.paragraphs.
"""
from docx import Document


def add_lines_to_cell(cell, lines: list[str]):
    """Add multiple lines as separate paragraphs to a cell."""
    # Clear default paragraph and set first line
    cell.paragraphs[0].text = lines[0]
    # Add remaining lines as new paragraphs
    for line in lines[1:]:
        cell.add_paragraph(line)


document = Document()
document.add_heading("UAT Test Document", level=1)

# --- Table 1: Single commands ---
document.add_heading("1. Single Commands", level=2)
table1 = document.add_table(rows=3, cols=2)
table1.style = "Table Grid"

# Row 0: display device
add_lines_to_cell(table1.rows[0].cells[0], [
    "<HUAWEI>display device",
    "<TUC-TYB91G01HWLEFC303-CPLEF03>",
    "<TUC-TYB91G01HWLEFC303-CPLEF04>",
])

# Row 1: display clock
add_lines_to_cell(table1.rows[1].cells[0], [
    "<HUAWEI>display clock",
    "<TUC-TYB91G01HWLEFC303-CPLEF03>",
])

# Row 2: display version
add_lines_to_cell(table1.rows[2].cells[0], [
    "<HUAWEI>display version",
    "<TUC-TYB91G01HWLEFC303-CPLEF03>",
])

# --- Table 2: Nested commands ---
document.add_heading("2. Nested Commands", level=2)
table2 = document.add_table(rows=3, cols=2)
table2.style = "Table Grid"

# Row 0: system-view → interface (full commands)
add_lines_to_cell(table2.rows[0].cells[0], [
    "<HUAWEI>system-view",
    "[~HUAWEI]interface GigabitEthernet0/0/29",
    "[~HUAWEI-GigabitEthernet0/0/29]display this",
    "[~HUAWEI-GigabitEthernet0/0/29]shutdown",
    "[~HUAWEI-GigabitEthernet0/0/29]undo shutdown",
    "[~HUAWEI-GigabitEthernet0/0/29]display this",
    "[~HUAWEI-GigabitEthernet0/0/29]quit",
    "[~HUAWEI]quit",
    "<TUC-TYB91G01HWLEFC303-CPLEF03>",
    "<TUC-TYB91G01HWLEFC303-CPLEF04>",
])

# Row 1: system-view → OSPF (abbreviated: sys, dis th, q)
add_lines_to_cell(table2.rows[1].cells[0], [
    "<HUAWEI>sys",
    "[~HUAWEI]ospf 1",
    "[~HUAWEI-ospf-1]dis th",
    "[~HUAWEI-ospf-1]network 172.16.0.0 0.0.255.255 area 0",
    "[~HUAWEI-ospf-1]silent-interface all",
    "[~HUAWEI-ospf-1]undo silent-interface all",
    "[~HUAWEI-ospf-1]q",
    "[~HUAWEI]q",
    "<TUC-TYB91G01HWLEFC303-CPLEF03>",
])

# Row 2: system-view → display commands (no sub-view)
add_lines_to_cell(table2.rows[2].cells[0], [
    "<HUAWEI>system-view",
    "[~HUAWEI]display current-configuration",
    "[~HUAWEI]display ip routing-table",
    "[~HUAWEI]display vlan",
    "[~HUAWEI]display interface brief",
    "[~HUAWEI]quit",
    "<TUC-TYB91G01HWLEFC303-CPLEF03>",
])

output_path = "test_document_v2.docx"
document.save(output_path)
print(f"Created {output_path}")