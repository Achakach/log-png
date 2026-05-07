from docx import Document

doc = Document()

# Reproduce checkblockabove_error with proper format for putpnginword
doc.add_heading('SSH Merge Error Case', level=1)

# This simulates what the template would look like:
# The cell contains only the command that executed on the target device
# If it failed, the word 'Error' would be in the output text
table = doc.add_table(rows=1, cols=1)
cell = table.rows[0].cells[0]
cell.add_paragraph("<TUCxx01>dis cur")
cell.add_paragraph("Error: Unrecognized command")
cell.add_paragraph("<TUCxx01>")

doc.save("ssh_error_checkblockabove.docx")
print("Created ssh_error_checkblockabove.docx")
