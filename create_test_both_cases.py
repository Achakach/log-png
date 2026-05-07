from docx import Document

# Create test DOCX with both cases
doc = Document()

# --- Case 1: Standalone display current-configuration (no SSH) ---
doc.add_heading('Case 1: Standalone dis cur (no SSH)', level=2)
table1 = doc.add_table(rows=1, cols=1)
cell1 = table1.rows[0].cells[0]
cell1.add_paragraph('<HUAWEI>display current-configuration')
cell1.add_paragraph('<TUCxx01>')
cell1.add_paragraph('<TUCxx02>')

# --- Case 2: SSH merge with username ---
doc.add_heading('Case 2: SSH merge with username', level=2)
table2 = doc.add_table(rows=1, cols=1)
cell2 = table2.rows[0].cells[0]
cell2.add_paragraph('<HUAWEI>display current-configuration')
cell2.add_paragraph('<HUAWEI>username kacha1')
cell2.add_paragraph('<TUCxx01>')
cell2.add_paragraph('<TUCxx02>')

# --- Case 3: DOCX without username, PNG has username (should NOT match) ---
doc.add_heading('Case 3: DOCX no username vs PNG with username', level=2)
table3 = doc.add_table(rows=1, cols=1)
cell3 = table3.rows[0].cells[0]
cell3.add_paragraph('<HUAWEI>display current-configuration')
cell3.add_paragraph('<TUCxx01>')

doc.save('test_both_cases.docx')
print('Created: test_both_cases.docx')
