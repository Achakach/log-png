from docx import Document

doc = Document(r"C:\Users\kacha\OneDrive\Desktop\test\comprehensive_test_cases.docx")

print(f"Total tables: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    cell_text = table.rows[1].cells[1].text.strip()[:200] if len(table.rows) > 1 else "N/A"
    print(f"\nTable {i+1}:")
    print(f"  {cell_text}")
