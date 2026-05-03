"""Generate a test case with actual [ERROR] PNGs to demonstrate the logic."""
from docx import Document
from docx.shared import Inches
import os

# Create test directory and mock PNG files
test_dir = 'test_error_scenario'
os.makedirs(test_dir, exist_ok=True)

# Create mock PNG files (just empty files with correct names)
# TUC01: device works, clock is error
# TUC02: device is error, clock works
# TUC03: both work
mock_pngs = [
    'TUC-TEST01 display device.png',          # clean
    'TUC-TEST01 display clock [error].png',    # error
    'TUC-TEST02 display device [error].png',   # error
    'TUC-TEST02 display clock.png',           # clean
    'TUC-TEST03 display device.png',          # clean
    'TUC-TEST03 display clock.png',           # clean
]

for png_name in mock_pngs:
    path = os.path.join(test_dir, png_name)
    if not os.path.exists(path):
        # Create a minimal valid PNG (1x1 transparent pixel)
        import struct
        png_data = b'\x89PNG\r\n\x1a\n' + struct.pack('>I', 13) + b'IHDR' + struct.pack('>IIBBBBB', 1, 1, 8, 6, 0, 0, 0) + b'\x00\x00\x00\x00\x00\x00\x00\x00' + struct.pack('>I', 0) + b'IEND' + b'\xae\x42\x60\x82'
        with open(path, 'wb') as f:
            f.write(png_data)
        print(f"Created: {png_name}")

# Create DOCX with Case 12 structure
doc = Document()
doc.add_heading('Case 12: [ERROR] PNG Test', level=1)
doc.add_paragraph('TUC01 supports display device (clock is error)')
doc.add_paragraph('TUC02 supports display clock (device is error)')
doc.add_paragraph('TUC03 supports both')

table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'

header_cells = table.rows[0].cells
header_cells[0].text = 'Case'
header_cells[1].text = 'Cell Content'

content_cells = table.rows[1].cells
content_cells[0].text = 'Case 12'

cell = content_cells[1]
cell.text = ''
cell.add_paragraph('CEx01')
cell.add_paragraph('<HUAWEI>display device')
cell.add_paragraph('CEx02')
cell.add_paragraph('<HUAWEI>display clock')
cell.add_paragraph('<TUC-TEST01>')
cell.add_paragraph('<TUC-TEST02>')
cell.add_paragraph('<TUC-TEST03>')

doc.save('test_error_case.docx')
print("\nGenerated: test_error_case.docx")
print("\nExpected results:")
print("  TUC01: display device (clean) -> INSERT")
print("         display clock [ERROR] -> SKIP")
print("  TUC02: display device [ERROR] -> SKIP")
print("         display clock (clean) -> INSERT")
print("  TUC03: display device (clean) -> INSERT")
print("         display clock (clean) -> SKIP (already inserted)")
