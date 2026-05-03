from docx import Document
import re

# Check the generated result file
doc = Document('comprehensive_test_cases_RESULT.docx')

print('=== Case 12 from comprehensive_test_cases_RESULT.docx ===\n')
table12 = doc.tables[11]  # 0-indexed, so 11 = Case 12

for ri, row in enumerate(table12.rows):
    for ci, cell in enumerate(row.cells):
        text = cell.text.strip()
        if text and len(text) > 5:
            img_count = len(cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'))
            print(f'Row {ri} Col {ci}: {img_count} image(s)')
            for pi, para in enumerate(cell.paragraphs):
                print(f'  Para {pi}: {para.text}')
            print()

# Also check the source comprehensive_test_cases.docx
print('\n=== Source comprehensive_test_cases.docx ===\n')
doc2 = Document('comprehensive_test_cases.docx')
table12_src = doc2.tables[11]

for ri, row in enumerate(table12_src.rows):
    for ci, cell in enumerate(row.cells):
        text = cell.text.strip()
        if text and len(text) > 5:
            print(f'Row {ri} Col {ci}:')
            for pi, para in enumerate(cell.paragraphs):
                print(f'  Para {pi}: {para.text}')
            print()
