"""Inspect comprehensive_test_cases_updated.docx to find commands."""
from docx import Document

doc = Document('comprehensive_test_cases_updated.docx')
print(f"Tables: {len(doc.tables)}")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti+1} ---")
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                lines = text.split('\n')
                if len(lines) <= 5:  # Only show small cells
                    print(f"  Row {ri+1}, Col {ci+1}: {lines}")
                else:
                    print(f"  Row {ri+1}, Col {ci+1}: ({len(lines)} lines)")
                    print(f"    First: {lines[0][:80]}")
                    print(f"    Last:  {lines[-1][:80]}")
