from docx import Document

doc = Document()

doc.add_heading('dis cur example!', level=1)

# Case 1: Simple - no username
doc.add_paragraph('Case 1: DOCX has no username command')
table1 = doc.add_table(rows=1, cols=1)
cell1 = table1.rows[0].cells[0]
cell1.add_paragraph('<HUAWEI>display current-configuration')
cell1.add_paragraph('<TUCxx01>')  # Node 1
cell1.add_paragraph('<TUCxx02>')  # Node 2

# Case 2: With username as metadata
doc.add_paragraph('Case 2: DOCX has username as part of matching')
table2 = doc.add_table(rows=1, cols=1)
cell2 = table2.rows[0].cells[0]
cell2.add_paragraph('<HUAWEI>display current-configuration')
cell2.add_paragraph('<HUAWEI>username kacha1')
cell2.add_paragraph('<TUCxx01>')  # Node 1
cell2.add_paragraph('<TUCxx02>')  # Node 2

doc.save('dis_cur_example.docx')
print('Created: dis_cur_example.docx')
