from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('UAT Test Document', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add some body text
p = doc.add_paragraph()
p.add_run('CloudEngine (Switch)\n').bold = True
p.add_run('Site Code: TYB DC61')

p = doc.add_paragraph()
p.add_run('Representative of TRUE: ').bold = True
p.add_run('Huawei Technologies Co., Ltd.')

doc.add_page_break()

# Section 1: Single Commands
doc.add_heading('1. Single Commands', level=1)

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Objective'
hdr_cells[1].text = 'To test that display device command works'
hdr_cells[2].text = 'To test that display version command works'

# Row 1
row_cells = table.rows[1].cells
row_cells[0].text = 'Test network diagram'
row_cells[1].text = 'none'
row_cells[2].text = 'none'

# Row 2
row_cells = table.rows[2].cells
row_cells[0].text = 'Preset condition'
row_cells[1].text = 'the switch is powered on and connected'
row_cells[2].text = 'the switch is powered on and connected'

# Row 3 - Procedure header
row_cells = table.rows[3].cells
row_cells[0].text = 'Test procedure'
row_cells[1].text = 'Test procedure'
row_cells[2].text = 'Expected result'

# Row 4 - Single command 1
row_cells = table.rows[4].cells
row_cells[0].text = '1. Run display device'
row_cells[1].text = '1. Run display device'
row_cells[2].text = 'Take the display device screenshot\n<HUAWEI>display device\n<TUC-TYB91G01HWLEFC303-CPLEF03>\n<TUC-TYB91G01HWLEFC303-CPLEF04>'

# Row 5 - Single command 2
row_cells = table.rows[5].cells
row_cells[0].text = '2. Run display version'
row_cells[1].text = '2. Run display version'
row_cells[2].text = '<HUAWEI>display version\n<TUC-TYB91G01HWLEFC303-CPLEF03>\n<TUC-TYB91G01HWLEFC303-CPLEF04>'

# Row 6 - Test description
row_cells = table.rows[6].cells
row_cells[0].text = 'Test description'
row_cells[1].text = 'The actual screenshot should show device information'
row_cells[2].text = 'The actual screenshot should show version information'

doc.add_paragraph()

# Section 2: Nested Commands
doc.add_heading('2. Nested Commands', level=1)

# 2.1 Nested command ex1
doc.add_heading('2.1 Configure Loopback Interface', level=2)

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Objective'
hdr_cells[1].text = 'To test that loopback interface configuration works'
hdr_cells[2].text = 'To test that loopback interface display works'

# Row 1
row_cells = table.rows[1].cells
row_cells[0].text = 'Test network diagram'
row_cells[1].text = 'none'
row_cells[2].text = 'none'

# Row 2
row_cells = table.rows[2].cells
row_cells[0].text = 'Preset condition'
row_cells[1].text = 'the switch is in system view'
row_cells[2].text = 'the switch is in system view'

# Row 3
row_cells = table.rows[3].cells
row_cells[0].text = 'Test procedure'
row_cells[1].text = 'Test procedure'
row_cells[2].text = 'Expected result'

# Row 4 - Nested ex1
row_cells = table.rows[4].cells
row_cells[0].text = '1. Configure loopback 123'
row_cells[1].text = '1. Configure loopback 123'
row_cells[2].text = 'Take the configuration screenshot\n<HUAWEI>system\n[~HUAWEI]interface loopback 123\n[*HUAWEI-Loopback123]commit\n[~HUAWEI-Loopback123]dis th\n#\ninterface Loopback123\n#\n[~HUAWEI-Loopback123]q\n[~HUAWEI]undo interface LoopBack 123\n[*HUAWEI]command\n[~HUAWEI]\n<TUC-TYB91G01HWLEFC303-CPLEF03>\n<TUC-TYB91G01HWLEFC303-CPLEF04>'

# Row 5 - Nested ex2
row_cells = table.rows[5].cells
row_cells[0].text = '2. Set CPU threshold'
row_cells[1].text = '2. Set CPU threshold'
row_cells[2].text = '<HUAWEI>system-view\n[HUAWEI]set cpu threshold 4\n<TUC-TYB91G01HWLEFC303-CPLEF03>\n<TUC-TYB91G01HWLEFC303-CPLEF04>'

# Row 6
row_cells = table.rows[6].cells
row_cells[0].text = 'Test description'
row_cells[1].text = 'The actual screenshot should show configuration'
row_cells[2].text = 'The actual screenshot should show CPU settings'

doc.add_paragraph()

# 2.2 Multi-command in one cell
doc.add_heading('2.2 Multiple Commands in One Cell', level=2)

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Objective'
hdr_cells[1].text = 'To test multiple command blocks'
hdr_cells[2].text = 'Expected result'

# Row 1
row_cells = table.rows[1].cells
row_cells[0].text = 'Test network diagram'
row_cells[1].text = 'none'
row_cells[2].text = 'none'

# Row 2
row_cells = table.rows[2].cells
row_cells[0].text = 'Preset condition'
row_cells[1].text = 'the switch is connected'
row_cells[2].text = 'the switch is connected'

# Row 3
row_cells = table.rows[3].cells
row_cells[0].text = 'Test procedure'
row_cells[1].text = 'Test procedure'
row_cells[2].text = 'Expected result'

# Row 4 - Multi command 1
row_cells = table.rows[4].cells
row_cells[0].text = '1. Command X'
row_cells[1].text = '1. Command X'
row_cells[2].text = 'Take screenshot X\n<HUAWEI>commandx1\n[HUAWEI]commandx2\n<device1>\npicx1\n<device2>\npicx2'

# Row 5 - Multi command 2
row_cells = table.rows[5].cells
row_cells[0].text = '2. Command Y'
row_cells[1].text = '2. Command Y'
row_cells[2].text = 'Take screenshot Y\n<HUAWEI>commandy1\n[HUAWEI]commandy2\n<device1>\npicy1\n<device2>\npicy2'

# Row 6
row_cells = table.rows[6].cells
row_cells[0].text = 'Test description'
row_cells[1].text = 'Multiple command blocks in same cell'
row_cells[2].text = 'Each block should have its own screenshot'

doc.add_paragraph()

# Section 3: Different commands for different NE
doc.add_heading('3. Different Commands for Different NE Models', level=1)

# 3.1 CPU display
doc.add_heading('3.1 Display CPU Usage', level=2)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Objective'
hdr_cells[1].text = 'To test CPU display on different models'
hdr_cells[2].text = 'Expected result'

# Row 1
row_cells = table.rows[1].cells
row_cells[0].text = 'Test network diagram'
row_cells[1].text = 'none'
row_cells[2].text = 'none'

# Row 2
row_cells = table.rows[2].cells
row_cells[0].text = 'Preset condition'
row_cells[1].text = 'NEs are powered on'
row_cells[2].text = 'NEs are powered on'

# Row 3
row_cells = table.rows[3].cells
row_cells[0].text = 'Test procedure'
row_cells[1].text = 'Test procedure'
row_cells[2].text = 'Expected result'

# Row 4 - Different commands
row_cells = table.rows[4].cells
row_cells[0].text = '1. Display CPU'
row_cells[1].text = '1. Display CPU'
row_cells[2].text = 'CE885:\n<HUAWEI>display cpu-usage\nCE6863E/CE16816:\n<HUAWEI>display cpu\n\n<device1>\n<device2>'

# Row 5
row_cells = table.rows[5].cells
row_cells[0].text = 'Test description'
row_cells[1].text = 'Different commands for different NE models'
row_cells[2].text = 'Each model uses its appropriate command'

doc.add_paragraph()

# 3.2 Display memory
doc.add_heading('3.2 Display Memory', level=2)

table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'

# Header
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Objective'
hdr_cells[1].text = 'To test memory display on different models'
hdr_cells[2].text = 'Expected result'

# Row 1
row_cells = table.rows[1].cells
row_cells[0].text = 'Test network diagram'
row_cells[1].text = 'none'
row_cells[2].text = 'none'

# Row 2
row_cells = table.rows[2].cells
row_cells[0].text = 'Preset condition'
row_cells[1].text = 'NEs are powered on'
row_cells[2].text = 'NEs are powered on'

# Row 3
row_cells = table.rows[3].cells
row_cells[0].text = 'Test procedure'
row_cells[1].text = 'Test procedure'
row_cells[2].text = 'Expected result'

# Row 4 - Different commands
row_cells = table.rows[4].cells
row_cells[0].text = '1. Display Memory'
row_cells[1].text = '1. Display Memory'
row_cells[2].text = 'CE885:\n<HUAWEI>display memory-usage\nCE6863E/CE16816:\n<HUAWEI>display memory\n\n<device1>\n<device2>'

# Row 5
row_cells = table.rows[5].cells
row_cells[0].text = 'Test description'
row_cells[1].text = 'Different memory commands for different NE models'
row_cells[2].text = 'Each model uses its appropriate command'

doc.save('test_document_comprehensive.docx')
print('Created test_document_comprehensive.docx')
